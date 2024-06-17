import os
import requests
import random
import pdfkit
from bs4 import BeautifulSoup
import concurrent.futures
from pathlib import Path
from docx import Document
from zipfile import ZipFile
from pywebio.input import input_group, textarea, input, NUMBER
from pywebio.output import put_file, put_success, put_error

def extract_delivery_notices(url, headers):
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
    except requests.RequestException as e:
        print(f"请求错误: {e}")
        return []

    soup = BeautifulSoup(response.text, 'html.parser')
    delivery_notices = soup.find_all('li')
    notices = []
    for notice in delivery_notices:
        a_tag = notice.find('a')
        if a_tag and 'href' in a_tag.attrs:
            title = a_tag.text.strip()
            link = a_tag['href']
            if link.startswith('/'):
                link = "http://www.gzcourt.gov.cn" + link
            notices.append((title, link, url))
    return notices

def fetch_notices(page, headers, keyword):
    base_url = "https://www.gzcourt.gov.cn/other/ck601/index{}.html"
    if page == 1:
        url = "https://www.gzcourt.gov.cn/other/ck601/index.html"
    else:
        url = base_url.format(page - 1)

    notices = extract_delivery_notices(url, headers)
    return [(page, notice) for notice in notices if keyword in notice[0]]

def save_page_as_pdf(url, filename):
    try:
        pdfkit.from_url(url, filename)
    except Exception as e:
        print(f"保存PDF时出错: {e}")

def process_notice(notice_info, doc, keyword, generated_files):
    page, (notice_title, notice_link, page_url) = notice_info
    doc.add_paragraph(f"在第 {page} 页找到了对应公告：")
    doc.add_paragraph("公告标题: " + notice_title)
    doc.add_paragraph("公告链接: " + notice_link)
    doc.add_paragraph("公告所在页面链接: " + page_url)

    pdf_page_name = f"{keyword}_page_{page}.pdf"
    save_page_as_pdf(page_url, pdf_page_name)
    generated_files.append(pdf_page_name)

    pdf_notice_name = f"{keyword}_notice_{page}.pdf"
    save_page_as_pdf(notice_link, pdf_notice_name)
    generated_files.append(pdf_notice_name)

def search_delivery_notices(doc, keyword, start_page, end_page, generated_files):
    user_agents = [
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36',
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
    ]
    headers = {
        'User-Agent': random.choice(user_agents),
        'Accept-Language': 'en-US,en;q=0.9'
    }

    with concurrent.futures.ThreadPoolExecutor() as executor:
        futures = [executor.submit(fetch_notices, page, headers, keyword) for page in range(start_page, end_page + 1)]
        for future in concurrent.futures.as_completed(futures):
            notices = future.result()
            for notice in notices:
                process_notice(notice, doc, keyword, generated_files)

def main():
    inputs = input_group("搜索送达公告", [
        textarea("请输入要搜索的案号（每行一个）：", name="keywords"),
        input("请输入起始页码：", name="start_page", type=NUMBER),
        input("请输入结束页码：", name="end_page", type=NUMBER)
    ])

    keywords = inputs['keywords'].strip().split()
    start_page = inputs['start_page']
    end_page = inputs['end_page']

    all_files = []

    for keyword in keywords:
        doc = Document()
        generated_files = []

        search_delivery_notices(doc, keyword, start_page, end_page, generated_files)

        if len(doc.paragraphs) == 0:
            doc.add_paragraph("没有找到与任何案号相关的送达公告。")

        doc_filename = f"delivery_notices_{keyword}.docx"
        doc.save(doc_filename)
        all_files.append(doc_filename)

        for file in generated_files:
            all_files.append(file)

        put_file(doc_filename, open(doc_filename, 'rb'), '下载生成的 DOCX 文件')

        for pdf_file in generated_files:
            if os.path.exists(pdf_file):  # 确保文件存在
                put_file(pdf_file, open(pdf_file, 'rb'), f'下载生成的 PDF 文件: {pdf_file}')
            else:
                print(f"文件 {pdf_file} 不存在，无法下载。")

    zip_filename = "all_generated_files.zip"
    with ZipFile(zip_filename, 'w') as zipf:
        for file in all_files:
            zipf.write(file, os.path.basename(file))

    put_file(zip_filename, open(zip_filename, 'rb'), '下载所有生成的文件 (ZIP)')
    put_success("所有文件已生成并打包，点击上方按钮下载。")

# Vercel serverless function entry point
def handler(request):
    main()

